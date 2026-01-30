"""
Document Import Script
Reads .docx files from raw/ folder and extracts structured text.
"""

import os
import json
from pathlib import Path
from docx import Document
from datetime import datetime

class DocumentImporter:
    def __init__(self, raw_folder='../raw', processed_folder='../processed'):
        self.raw_folder = Path(__file__).parent / raw_folder
        self.processed_folder = Path(__file__).parent / processed_folder
        
        # Create folders if they don't exist
        self.raw_folder.mkdir(parents=True, exist_ok=True)
        self.processed_folder.mkdir(parents=True, exist_ok=True)
    
    def extract_text_from_docx(self, docx_path):
        """
        Extract text from a .docx file with structure preservation.
        Returns a list of sections with text and metadata.
        """
        doc = Document(docx_path)
        sections = []
        current_section = {
            'heading': 'Introduction',
            'level': 0,
            'content': []
        }
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            if not text:
                continue
            
            # Check if it's a heading
            if para.style.name.startswith('Heading'):
                # Save previous section if it has content
                if current_section['content']:
                    current_section['text'] = '\n'.join(current_section['content'])
                    sections.append(current_section.copy())
                
                # Start new section
                level = int(para.style.name.replace('Heading ', '')) if para.style.name != 'Heading' else 1
                current_section = {
                    'heading': text,
                    'level': level,
                    'content': []
                }
            else:
                # Regular paragraph
                current_section['content'].append(text)
        
        # Add the last section
        if current_section['content']:
            current_section['text'] = '\n'.join(current_section['content'])
            sections.append(current_section)
        
        # If no sections were created (no headings), create one section with all content
        if not sections:
            all_text = '\n\n'.join([para.text.strip() for para in doc.paragraphs if para.text.strip()])
            sections.append({
                'heading': 'Document Content',
                'level': 1,
                'text': all_text
            })
        
        return sections
    
    def process_document(self, filename):
        """
        Process a single .docx file and save as JSON.
        """
        docx_path = self.raw_folder / filename
        
        if not docx_path.exists():
            print(f"Error: File {filename} not found in {self.raw_folder}")
            return None
        
        print(f"Processing: {filename}")
        
        try:
            sections = self.extract_text_from_docx(docx_path)
            
            # Create document metadata
            doc_data = {
                'source_file': filename,
                'processed_date': datetime.now().isoformat(),
                'num_sections': len(sections),
                'sections': sections
            }
            
            # Save as JSON
            output_filename = f"{docx_path.stem}.json"
            output_path = self.processed_folder / output_filename
            
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(doc_data, f, indent=2, ensure_ascii=False)
            
            print(f"✓ Saved to: {output_path}")
            print(f"  - {len(sections)} sections extracted")
            
            return doc_data
            
        except Exception as e:
            print(f"Error processing {filename}: {str(e)}")
            return None
    
    def process_all_documents(self):
        """
        Process all .docx files in the raw folder.
        """
        docx_files = list(self.raw_folder.glob('*.docx'))
        
        if not docx_files:
            print(f"No .docx files found in {self.raw_folder}")
            print("\nTo use this script:")
            print("1. Export your OneNote notebooks as .docx files")
            print(f"2. Place them in: {self.raw_folder.absolute()}")
            print("3. Run this script again")
            return []
        
        print(f"Found {len(docx_files)} document(s) to process\n")
        
        results = []
        for docx_file in docx_files:
            result = self.process_document(docx_file.name)
            if result:
                results.append(result)
        
        print(f"\n✓ Successfully processed {len(results)} document(s)")
        return results


if __name__ == '__main__':
    importer = DocumentImporter()
    importer.process_all_documents()
